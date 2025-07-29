import os
import io
from typing import Optional, Dict, Any, List
from dataclasses import dataclass
import PyPDF2
import docx
import markdown
from pathlib import Path


@dataclass
class DocumentMetadata:
    """Metadata about a loaded document."""
    filename: str
    file_type: str
    file_size: int
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    encoding: Optional[str] = None


class DocumentLoader:
    """
    Universal document loader supporting multiple formats.

    Supported formats:
    - PDF (.pdf)
    - Word documents (.docx)
    - Markdown (.md)
    - Text files (.txt)
    - Raw text input
    """

    def __init__(self):
        self.supported_extensions = {
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.md': self._load_markdown,
            '.txt': self._load_text,
        }

    def load_document(self, source, source_type: str = "auto") -> tuple[str, DocumentMetadata]:
        """
        Load document from various sources.

        Args:
            source: File path, file-like object, or raw text
            source_type: "file", "upload", "text", or "auto"

        Returns:
            Tuple of (document_text, metadata)
        """
        if source_type == "text" or isinstance(source, str) and not os.path.exists(source):
            return self._load_raw_text(source)

        if source_type == "upload" or hasattr(source, 'read'):
            return self._load_uploaded_file(source)

        if source_type == "file" or (isinstance(source, str) and os.path.exists(source)):
            return self._load_file(source)

        # Auto-detect
        if isinstance(source, str):
            if os.path.exists(source):
                return self._load_file(source)
            else:
                return self._load_raw_text(source)
        else:
            return self._load_uploaded_file(source)

    def _load_file(self, file_path: str) -> tuple[str, DocumentMetadata]:
        """Load document from file path."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file format: {extension}")

        # Load content using appropriate method
        content = self.supported_extensions[extension](file_path)

        # Create metadata
        metadata = DocumentMetadata(
            filename=path.name,
            file_type=extension,
            file_size=path.stat().st_size,
            character_count=len(content),
            word_count=len(content.split())
        )

        return content, metadata

    def _load_uploaded_file(self, uploaded_file) -> tuple[str, DocumentMetadata]:
        """Load document from uploaded file object (Streamlit)."""
        filename = getattr(uploaded_file, 'name', 'uploaded_file')
        file_size = getattr(uploaded_file, 'size', 0)

        # Get file extension
        extension = Path(filename).suffix.lower()

        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file format: {extension}")

        # Read file content
        file_content = uploaded_file.read()

        # Reset file pointer if possible
        if hasattr(uploaded_file, 'seek'):
            uploaded_file.seek(0)

        # Load content based on file type
        if extension == '.pdf':
            content = self._load_pdf_from_bytes(file_content)
        elif extension == '.docx':
            content = self._load_docx_from_bytes(file_content)
        elif extension in ['.txt', '.md']:
            content = file_content.decode('utf-8')
        else:
            raise ValueError(f"Cannot process uploaded file of type: {extension}")

        # Create metadata
        metadata = DocumentMetadata(
            filename=filename,
            file_type=extension,
            file_size=file_size,
            character_count=len(content),
            word_count=len(content.split())
        )

        return content, metadata

    def _load_raw_text(self, text: str) -> tuple[str, DocumentMetadata]:
        """Load raw text input."""
        metadata = DocumentMetadata(
            filename="raw_text_input",
            file_type="text",
            file_size=len(text.encode('utf-8')),
            character_count=len(text),
            word_count=len(text.split())
        )

        return text, metadata

    def _load_pdf(self, file_path: str) -> str:
        """Load text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                return self._load_pdf_from_bytes(file.read())
        except Exception as e:
            raise Exception(f"Error loading PDF {file_path}: {str(e)}")

    def _load_pdf_from_bytes(self, pdf_bytes: bytes) -> str:
        """Load text from PDF bytes."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
            text_parts = []

            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)

            return '\n\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")

    def _load_docx(self, file_path: str) -> str:
        """Load text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            return '\n\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Error loading DOCX {file_path}: {str(e)}")

    def _load_docx_from_bytes(self, docx_bytes: bytes) -> str:
        """Load text from DOCX bytes."""
        try:
            doc = docx.Document(io.BytesIO(docx_bytes))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            return '\n\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")

    def _load_markdown(self, file_path: str) -> str:
        """Load text from Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()

            # Convert markdown to plain text (optional: keep markdown formatting)
            # html = markdown.markdown(md_content)
            # You could use BeautifulSoup to extract text from HTML if needed

            return md_content
        except Exception as e:
            raise Exception(f"Error loading Markdown {file_path}: {str(e)}")

    def _load_text(self, file_path: str) -> str:
        """Load text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise Exception(f"Could not decode text file {file_path}")
        except Exception as e:
            raise Exception(f"Error loading text file {file_path}: {str(e)}")

    def get_sample_documents(self) -> Dict[str, str]:
        """Get sample documents for testing."""
        return {
            "UAV Research Paper": """
Energy-Efficient Inference on the Edge Exploiting TinyML Capabilities for UAVs

Abstract: In recent years, the proliferation of unmanned aerial vehicles (UAVs) has increased dramatically. UAVs can accomplish complex or dangerous tasks in a reliable and cost-effective way but are still limited by power consumption problems, which pose serious constraints on the flight duration and completion of energy-demanding tasks. The possibility of providing UAVs with advanced decision-making capabilities in an energy-effective way would be extremely beneficial.

In this paper, we propose a practical solution to this problem that exploits deep learning on the edge. The developed system integrates an OpenMV microcontroller into a DJI Tello Micro Aerial Vehicle (MAV). The microcontroller hosts a set of machine learning-enabled inference tools that cooperate to control the navigation of the drone and complete a given mission objective.

Introduction

Drones, in the form of both Remotely Piloted Aerial Systems (RPAS) and unmanned aerial vehicles (UAV), are increasingly being used to revolutionize many existing applications. The Internet of Things (IoT) is becoming more ubiquitous every day, thanks to the widespread adoption and integration of mobile robots into IoT ecosystems.

As the world becomes more dependent on technology, there is a growing need for autonomous systems that support the activities and mitigate the risks for human operators. In this context, UAVs are becoming increasingly popular in a range of civil and military applications such as smart agriculture, defense, construction site monitoring, and environmental monitoring.

These aerial vehicles are subject to numerous limitations such as safety, energy, weight, and space requirements. Electrically powered UAVs, which represent the majority of micro aerial vehicles, show a severe limitation in the duration of batteries, which are necessarily small due to design constraints.

Recent advances in embedded systems through IoT devices could open new and interesting possibilities in this domain. Edge computing brings new insights into existing IoT environments by solving many critical challenges. Deep learning (DL) at the edge presents significant advantages with respect to its distributed counterpart.

Another recent trend refers to the possibility of shifting the ML inference peripherally by exploiting new classes of microcontrollers, thus generating the notion of Tiny Machine Learning (TinyML). TinyML aims to bring ML inference into devices characterized by a very low power consumption.

Building upon the above technological trends, the integration of state-of-the-art ultra-low power embedded devices into UAVs could provide energy-aware solutions to embed an increasing amount of autonomy and intelligence into the drone.
            """,

            "Short Technical Document": """
Machine Learning at the Edge: A Comprehensive Overview

Introduction
Edge computing has emerged as a paradigm shift in how we process and analyze data. Unlike traditional cloud-based approaches, edge computing brings computation closer to the data source, reducing latency and improving real-time decision making.

Key Benefits
1. Reduced Latency: Processing data locally eliminates network delays
2. Enhanced Privacy: Sensitive data remains on local devices
3. Improved Reliability: Less dependence on network connectivity
4. Cost Optimization: Reduced bandwidth and cloud computing costs

Challenges
Despite its advantages, edge computing faces several challenges including limited computational resources, power constraints, and the need for efficient algorithms that can run on resource-constrained devices.

Conclusion
The future of edge computing looks promising with advances in hardware miniaturization and algorithm optimization. As IoT devices become more prevalent, edge computing will play a crucial role in enabling intelligent, responsive systems.
            """,

            "Business Report": """
Quarterly Business Review - Q3 2024

Executive Summary
Our organization has achieved significant milestones in Q3 2024, with revenue growth of 15% compared to the previous quarter. Key performance indicators show positive trends across all major business units.

Financial Performance
Total revenue reached $2.3 million, representing a 15% quarter-over-quarter increase. Profit margins improved by 3 percentage points due to operational efficiency improvements and cost optimization initiatives.

Market Analysis
The market continues to show strong demand for our core products. Customer satisfaction scores improved to 4.2/5.0, and customer retention rates reached 87%, the highest in company history.

Operational Highlights
- Launched three new product features
- Expanded team by 12 new employees
- Implemented new customer support system
- Achieved 99.8% system uptime

Future Outlook
Looking ahead to Q4, we anticipate continued growth driven by seasonal demand and new product launches. Our pipeline remains strong with several enterprise deals in advanced stages of negotiation.
            """
        }

    def validate_document(self, text: str, min_length: int = 100) -> tuple[bool, List[str]]:
        """
        Validate document content.

        Args:
            text: Document text to validate
            min_length: Minimum character length

        Returns:
            Tuple of (is_valid, list_of_issues)
        """
        issues = []

        if not text or not text.strip():
            issues.append("Document is empty")

        if len(text) < min_length:
            issues.append(f"Document too short (minimum {min_length} characters)")

        # Check for common issues
        if text.count('\x00') > 0:
            issues.append("Document contains null characters")

        # Check character distribution (detect garbled text)
        if len(text) > 100:
            alpha_ratio = sum(c.isalpha() for c in text) / len(text)
            if alpha_ratio < 0.3:
                issues.append("Document may contain garbled or non-text content")

        return len(issues) == 0, issues


# Convenience function for easy importing
def load_document(source, source_type: str = "auto") -> tuple[str, DocumentMetadata]:
    """
    Convenience function to load a document.

    Usage:
        text, metadata = load_document("document.pdf")
        text, metadata = load_document(uploaded_file, "upload")
        text, metadata = load_document("Raw text content", "text")
    """
    loader = DocumentLoader()
    return loader.load_document(source, source_type)


if __name__ == "__main__":
    # Example usage
    loader = DocumentLoader()

    # Test with sample documents
    samples = loader.get_sample_documents()

    for name, content in samples.items():
        text, metadata = loader.load_document(content, "text")
        print(f"\n{name}:")
        print(f"  Characters: {metadata.character_count}")
        print(f"  Words: {metadata.word_count}")
        print(f"  Preview: {text[:100]}...")

        # Validate
        is_valid, issues = loader.validate_document(text)
        print(f"  Valid: {is_valid}")
        if issues:
            print(f"  Issues: {', '.join(issues)}")
